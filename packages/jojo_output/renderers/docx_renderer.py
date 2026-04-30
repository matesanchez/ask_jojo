"""DOCX renderer — wraps the docx skill behind a typed spec.

PLAN.md Section 6 Phase 5: "Existing JoJo Bot docx/pptx/pdf/xlsx skills
are leveraged, not re-implemented." The skill at
``/sessions/.../skills/docx/SKILL.md`` defines the high-quality
pattern for producing professional Word documents; this renderer is
the typed-spec gateway that the synthesis path calls.

The spec is the *shape* of the document; the rendering uses
``python-docx`` to materialize it. For complex documents (TOC, page
numbers, multi-column layouts), the synthesis path can return a
richer spec; this renderer covers the common case (memo, report,
letter).

Public API:

- ``DocxSection`` / ``DocxTable`` / ``DocxSpec`` -- pydantic models.
- ``render_docx(spec, out_path)`` -- write a .docx file; return path.
"""

from __future__ import annotations

from pathlib import Path
from typing import Literal

from pydantic import BaseModel, Field


class DocxTable(BaseModel):
    """One table inside a section."""

    columns: list[str] = Field(..., min_length=1)
    rows: list[list[str]] = Field(default_factory=list)
    caption: str = ""


class DocxSection(BaseModel):
    """One section of the document."""

    heading: str = ""
    heading_level: int = Field(1, ge=1, le=4)
    body: str = ""
    bullets: list[str] = Field(default_factory=list)
    tables: list[DocxTable] = Field(default_factory=list)


class DocxSpec(BaseModel):
    """Full document spec.

    Attributes:
        title: document title; rendered as the topmost heading.
        subtitle: optional subtitle/byline.
        author: name on the cover (e.g. "JoJo Bot").
        style: ``memo`` (single-page memo formatting), ``report``
            (multi-page with TOC), or ``letter`` (formal letterhead).
        sections: ordered sections.
    """

    title: str = ""
    subtitle: str = ""
    author: str = "JoJo Bot"
    style: Literal["memo", "report", "letter"] = "report"
    sections: list[DocxSection] = Field(default_factory=list)


def render_docx(spec: DocxSpec, out_path: Path | str) -> Path:
    """Render a .docx file. Requires ``python-docx``.

    Returns the absolute path to the written file. Raises
    ``RuntimeError`` if python-docx isn't installed (with a clear
    install hint).
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for docx rendering; install via "
            "pip install -e \".[output]\""
        ) from e

    doc = Document()

    if spec.title:
        h = doc.add_heading(spec.title, level=0)
        for run in h.runs:
            run.font.size = Pt(20)
    if spec.subtitle:
        sub = doc.add_paragraph(spec.subtitle)
        sub.runs[0].italic = True
    if spec.author and spec.style != "memo":
        doc.add_paragraph(f"by {spec.author}")

    for section in spec.sections:
        if section.heading:
            doc.add_heading(section.heading, level=section.heading_level)
        if section.body:
            for para in section.body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        if section.bullets:
            for b in section.bullets:
                doc.add_paragraph(b, style="List Bullet")
        for tbl_spec in section.tables:
            tbl = doc.add_table(rows=1 + len(tbl_spec.rows), cols=len(tbl_spec.columns))
            tbl.style = "Light Grid Accent 1"
            # Header row.
            for i, col in enumerate(tbl_spec.columns):
                cell = tbl.rows[0].cells[i]
                cell.text = col
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            for r, row in enumerate(tbl_spec.rows, start=1):
                for c, val in enumerate(row[: len(tbl_spec.columns)]):
                    tbl.rows[r].cells[c].text = str(val) if val is not None else ""
            if tbl_spec.caption:
                cap = doc.add_paragraph(tbl_spec.caption)
                cap.runs[0].italic = True

    out = Path(out_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
